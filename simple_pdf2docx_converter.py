#!/usr/bin/env python3
"""
Simplified PDF to DOCX Converter with Template Support

This module provides functionality to convert PDF files to DOCX format with
support for custom templates (.dotx) and basic comparison capabilities.
"""

import os
import sys
import shutil
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
import logging
from dataclasses import dataclass

# Core conversion libraries
import fitz  # PyMuPDF
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Image processing
import cv2
import numpy as np
from PIL import Image

# CLI and utilities
import click
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table
from rich.panel import Panel

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

console = Console()


@dataclass
class ConversionConfig:
    """Configuration for PDF to DOCX conversion"""
    pdf_path: str
    output_path: str
    template_path: Optional[str] = None
    font_paths: Optional[List[str]] = None
    multi_processing: bool = True
    cpu_count: int = 4
    start_page: int = 0
    end_page: Optional[int] = None
    pages: Optional[List[int]] = None
    password: Optional[str] = None


class FontManager:
    """Basic font manager (simplified version)"""
    
    def __init__(self):
        self.installed_fonts = {}
        self.font_mappings = {}
    
    def install_fonts_from_directory(self, font_dir: str) -> int:
        """Install all fonts from a directory (basic implementation)"""
        font_dir = Path(font_dir)
        if not font_dir.exists():
            logger.error(f"Font directory not found: {font_dir}")
            return 0
        
        installed_count = 0
        for font_file in font_dir.glob("*"):
            if font_file.suffix.lower() in ['.ttf', '.otf']:
                # For now, just record the font files
                self.installed_fonts[font_file.stem] = str(font_file)
                installed_count += 1
                logger.info(f"Registered font: {font_file.name}")
        
        return installed_count
    
    def get_available_fonts(self) -> Dict[str, str]:
        """Get list of available fonts"""
        return self.installed_fonts.copy()


class TemplateManager:
    """Manages DOCX template operations"""
    
    def __init__(self):
        self.template_cache = {}
    
    def load_template(self, template_path: str) -> Optional[Document]:
        """Load a DOCX template file"""
        try:
            template_path = Path(template_path)
            if not template_path.exists():
                logger.error(f"Template file not found: {template_path}")
                return None
            
            if template_path.suffix.lower() not in ['.dotx', '.docx']:
                logger.error(f"Unsupported template format: {template_path.suffix}")
                return None
            
            # Load template
            template_doc = Document(str(template_path))
            self.template_cache[str(template_path)] = template_doc
            logger.info(f"Loaded template: {template_path.name}")
            return template_doc
            
        except Exception as e:
            logger.error(f"Error loading template {template_path}: {e}")
            return None
    
    def apply_template_styles(self, document: Document, template: Document) -> bool:
        """Apply template styles to a document"""
        try:
            # Basic style copying - simplified version
            template_styles = list(template.styles)
            document_style_names = [s.name for s in document.styles]
            
            styles_applied = 0
            for style in template_styles:
                if style.name not in document_style_names:
                    try:
                        # Try to add the style
                        new_style = document.styles.add_style(style.name, style.type)
                        if hasattr(style, 'font') and hasattr(new_style, 'font'):
                            if style.font.name:
                                new_style.font.name = style.font.name
                            if style.font.size:
                                new_style.font.size = style.font.size
                            if style.font.bold is not None:
                                new_style.font.bold = style.font.bold
                            if style.font.italic is not None:
                                new_style.font.italic = style.font.italic
                        styles_applied += 1
                    except Exception as style_error:
                        logger.debug(f"Could not apply style {style.name}: {style_error}")
                        continue
            
            logger.info(f"Applied {styles_applied} template styles successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error applying template styles: {e}")
            return False


class DocumentComparator:
    """Compares PDF and DOCX documents for accuracy"""
    
    def __init__(self):
        self.comparison_results = {}
    
    def extract_pdf_text(self, pdf_path: str) -> List[str]:
        """Extract text from PDF pages"""
        try:
            doc = fitz.open(pdf_path)
            pages_text = []
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                pages_text.append(text)
            
            doc.close()
            return pages_text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return []
    
    def extract_docx_text(self, docx_path: str) -> str:
        """Extract text from DOCX document"""
        try:
            doc = Document(docx_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def compare_documents(self, pdf_path: str, docx_path: str) -> Dict[str, Any]:
        """Compare PDF and DOCX documents"""
        try:
            pdf_text_pages = self.extract_pdf_text(pdf_path)
            docx_text = self.extract_docx_text(docx_path)
            
            pdf_text_combined = '\n'.join(pdf_text_pages)
            
            # Basic text comparison
            pdf_words = pdf_text_combined.split()
            docx_words = docx_text.split()
            
            # Calculate similarity metrics
            common_words = set(pdf_words) & set(docx_words)
            total_words = set(pdf_words) | set(docx_words)
            
            similarity = len(common_words) / len(total_words) if total_words else 0
            
            results = {
                'similarity_score': similarity,
                'pdf_word_count': len(pdf_words),
                'docx_word_count': len(docx_words),
                'common_words': len(common_words),
                'pdf_pages': len(pdf_text_pages),
                'status': 'success'
            }
            
            return results
            
        except Exception as e:
            logger.error(f"Error comparing documents: {e}")
            return {'status': 'error', 'error': str(e)}


class PDF2DOCXConverter:
    """Main converter class with template support"""
    
    def __init__(self):
        self.font_manager = FontManager()
        self.template_manager = TemplateManager()
        self.comparator = DocumentComparator()
    
    def setup_fonts(self, font_paths: List[str]) -> int:
        """Setup fonts from provided paths"""
        installed_count = 0
        
        for font_path in font_paths:
            font_path = Path(font_path)
            
            if font_path.is_dir():
                installed_count += self.font_manager.install_fonts_from_directory(str(font_path))
            elif font_path.is_file():
                # For now, just register the font file
                self.font_manager.installed_fonts[font_path.stem] = str(font_path)
                installed_count += 1
        
        return installed_count
    
    def convert_pdf_to_docx(self, config: ConversionConfig) -> Dict[str, Any]:
        """Convert PDF to DOCX with template support"""
        try:
            # Validate input file
            pdf_path = Path(config.pdf_path)
            if not pdf_path.exists():
                return {'status': 'error', 'error': f'PDF file not found: {pdf_path}'}
            
            # Setup output path
            output_path = Path(config.output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Install fonts if provided
            if config.font_paths:
                font_count = self.setup_fonts(config.font_paths)
                logger.info(f"Registered {font_count} fonts")
            
            # Load template if provided
            template_doc = None
            if config.template_path:
                template_doc = self.template_manager.load_template(config.template_path)
            
            # Perform PDF to DOCX conversion
            console.print("[bold green]Converting PDF to DOCX...[/bold green]")
            cv = Converter(str(pdf_path), password=config.password)
            
            # Set conversion parameters
            convert_params = {
                'docx_file': str(output_path),
                'multi_processing': config.multi_processing,
                'cpu_count': config.cpu_count
            }
            
            if config.pages:
                convert_params['pages'] = config.pages
            elif config.start_page or config.end_page:
                if config.start_page:
                    convert_params['start'] = config.start_page
                if config.end_page:
                    convert_params['end'] = config.end_page
            
            # Convert
            cv.convert(**convert_params)
            cv.close()
            
            # Verify output file was created
            if not output_path.exists():
                return {'status': 'error', 'error': f'Output file was not created: {output_path}'}
            
            # Apply template if provided
            if template_doc and output_path.exists():
                console.print("[bold blue]Applying template...[/bold blue]")
                doc = Document(str(output_path))
                self.template_manager.apply_template_styles(doc, template_doc)
                doc.save(str(output_path))
            
            # Compare documents
            console.print("[bold yellow]Comparing documents...[/bold yellow]")
            comparison_results = self.comparator.compare_documents(
                str(pdf_path), str(output_path)
            )
            
            return {
                'status': 'success',
                'output_path': str(output_path),
                'comparison': comparison_results,
                'fonts_installed': len(self.font_manager.get_available_fonts()),
                'template_applied': template_doc is not None
            }
            
        except Exception as e:
            logger.error(f"Conversion error: {e}")
            return {'status': 'error', 'error': str(e)}


@click.command()
@click.argument('pdf_path', type=click.Path(exists=True))
@click.argument('output_path', type=click.Path())
@click.option('--template', '-t', type=click.Path(exists=True), help='DOCX template file (.dotx)')
@click.option('--fonts', '-f', multiple=True, help='Font files or directories')
@click.option('--start-page', type=int, help='Start page (0-indexed)')
@click.option('--end-page', type=int, help='End page (exclusive)')
@click.option('--pages', help='Specific pages (comma-separated, 0-indexed)')
@click.option('--password', help='PDF password if encrypted')
@click.option('--no-multiprocessing', is_flag=True, help='Disable multiprocessing')
@click.option('--cpu-count', type=int, default=4, help='Number of CPUs to use')
def main(pdf_path, output_path, template, fonts, start_page, end_page, pages, 
         password, no_multiprocessing, cpu_count):
    """Convert PDF to DOCX with template support"""
    
    console.print(Panel.fit(
        "[bold blue]PDF to DOCX Converter[/bold blue]\n"
        "With Template Support",
        border_style="blue"
    ))
    
    # Parse pages if provided
    pages_list = None
    if pages:
        try:
            pages_list = [int(p.strip()) for p in pages.split(',')]
        except ValueError:
            console.print("[red]Error: Invalid pages format. Use comma-separated integers.[/red]")
            sys.exit(1)
    
    # Create configuration
    config = ConversionConfig(
        pdf_path=pdf_path,
        output_path=output_path,
        template_path=template,
        font_paths=list(fonts) if fonts else None,
        multi_processing=not no_multiprocessing,
        cpu_count=cpu_count,
        start_page=start_page or 0,
        end_page=end_page,
        pages=pages_list,
        password=password
    )
    
    # Create converter and run conversion
    converter = PDF2DOCXConverter()
    result = converter.convert_pdf_to_docx(config)
    
    # Display results
    if result['status'] == 'success':
        table = Table(title="Conversion Results")
        table.add_column("Metric", style="cyan")
        table.add_column("Value", style="green")
        
        table.add_row("Status", "✅ Success")
        table.add_row("Output File", result['output_path'])
        table.add_row("Fonts Registered", str(result['fonts_installed']))
        table.add_row("Template Applied", "✅ Yes" if result['template_applied'] else "❌ No")
        
        if 'comparison' in result and result['comparison']['status'] == 'success':
            comp = result['comparison']
            table.add_row("Similarity Score", f"{comp['similarity_score']:.2%}")
            table.add_row("PDF Words", str(comp['pdf_word_count']))
            table.add_row("DOCX Words", str(comp['docx_word_count']))
            table.add_row("PDF Pages", str(comp['pdf_pages']))
        
        console.print(table)
        console.print(f"\n[green]✅ Conversion completed successfully![/green]")
        console.print(f"[blue]Output saved to: {result['output_path']}[/blue]")
        
    else:
        console.print(f"[red]❌ Conversion failed: {result['error']}[/red]")
        sys.exit(1)


if __name__ == "__main__":
    main()
