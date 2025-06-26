#!/usr/bin/env python3
"""
PDF to DOCX Comparison and Adjustment Tool

This tool compares the original PDF with the converted DOCX and makes
adjustments to improve accuracy and formatting.
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import difflib
import re
import logging
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import json

# Set up logging
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

class PDFDOCXComparator:
    """Compare PDF and DOCX files and suggest/apply improvements"""
    
    def __init__(self):
        self.pdf_content = {}
        self.docx_content = {}
        self.comparison_results = {}
        self.adjustments_made = []
    
    def extract_pdf_content(self, pdf_path: str) -> Dict:
        """Extract detailed content from PDF for comparison"""
        try:
            pdf_doc = fitz.open(pdf_path)
            content = {
                'pages': [],
                'total_pages': len(pdf_doc),
                'text_blocks': [],
                'images': [],
                'fonts': set(),
                'colors': set()
            }
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                page_content = {
                    'page_num': page_num,
                    'text_blocks': [],
                    'images': [],
                    'raw_text': page.get_text(),
                    'bbox': page.rect
                }
                
                # Extract text with detailed formatting
                text_dict = page.get_text("dict")
                for block in text_dict["blocks"]:
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            line_text = ""
                            line_fonts = []
                            line_sizes = []
                            
                            for span in line["spans"]:
                                if span["text"].strip():
                                    line_text += span["text"]
                                    line_fonts.append(span["font"])
                                    line_sizes.append(span["size"])
                                    content['fonts'].add(span["font"])
                                    content['colors'].add(span["color"])
                            
                            if line_text.strip():
                                page_content['text_blocks'].append({
                                    'text': line_text,
                                    'bbox': line["bbox"],
                                    'fonts': line_fonts,
                                    'sizes': line_sizes
                                })
                
                # Extract images
                for img in page.get_images():
                    page_content['images'].append({
                        'xref': img[0],
                        'bbox': page.get_image_rects(img[0])
                    })
                
                content['pages'].append(page_content)
                content['text_blocks'].extend(page_content['text_blocks'])
            
            pdf_doc.close()
            self.pdf_content = content
            logger.info(f"Extracted PDF content: {len(content['text_blocks'])} text blocks, {len(content['images'])} images")
            return content
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return {}
    
    def extract_docx_content(self, docx_path: str) -> Dict:
        """Extract content from DOCX for comparison"""
        try:
            doc = Document(docx_path)
            content = {
                'paragraphs': [],
                'text_blocks': [],
                'images': [],
                'fonts': set(),
                'styles': []
            }
            
            for para in doc.paragraphs:
                if para.text.strip():
                    para_info = {
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'alignment': para.alignment,
                        'runs': []
                    }
                    
                    for run in para.runs:
                        if run.text.strip():
                            run_info = {
                                'text': run.text,
                                'font_name': run.font.name,
                                'font_size': run.font.size.pt if run.font.size else None,
                                'bold': run.font.bold,
                                'italic': run.font.italic,
                                'color': str(run.font.color.rgb) if run.font.color.rgb else None
                            }
                            para_info['runs'].append(run_info)
                            if run.font.name:
                                content['fonts'].add(run.font.name)
                    
                    content['paragraphs'].append(para_info)
                    content['text_blocks'].append(para.text)
            
            # Count images in document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    content['images'].append({'type': 'embedded'})
            
            self.docx_content = content
            logger.info(f"Extracted DOCX content: {len(content['paragraphs'])} paragraphs, {len(content['images'])} images")
            return content
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            return {}
    
    def compare_text_content(self) -> Dict:
        """Compare text content between PDF and DOCX"""
        if not self.pdf_content or not self.docx_content:
            return {}
        
        # Get all text from both sources
        pdf_text = "\n".join([block['text'] for block in self.pdf_content['text_blocks']])
        docx_text = "\n".join(self.docx_content['text_blocks'])
        
        # Clean text for comparison
        pdf_clean = re.sub(r'\s+', ' ', pdf_text).strip()
        docx_clean = re.sub(r'\s+', ' ', docx_text).strip()
        
        # Calculate similarity
        similarity = difflib.SequenceMatcher(None, pdf_clean, docx_clean).ratio()
        
        # Find differences
        differ = difflib.unified_diff(
            pdf_clean.splitlines(keepends=True),
            docx_clean.splitlines(keepends=True),
            fromfile='PDF',
            tofile='DOCX',
            lineterm=''
        )
        
        differences = list(differ)
        
        comparison = {
            'similarity_ratio': similarity,
            'pdf_text_length': len(pdf_text),
            'docx_text_length': len(docx_text),
            'missing_text': [],
            'extra_text': [],
            'differences': differences,
            'word_count_pdf': len(pdf_text.split()),
            'word_count_docx': len(docx_text.split())
        }
        
        # Identify missing and extra text
        pdf_words = set(pdf_clean.split())
        docx_words = set(docx_clean.split())
        
        comparison['missing_words'] = list(pdf_words - docx_words)
        comparison['extra_words'] = list(docx_words - pdf_words)
        
        return comparison
    
    def compare_formatting(self) -> Dict:
        """Compare formatting between PDF and DOCX"""
        formatting_comparison = {
            'fonts': {
                'pdf_fonts': list(self.pdf_content.get('fonts', [])),
                'docx_fonts': list(self.docx_content.get('fonts', [])),
                'missing_fonts': [],
                'font_mapping': {}
            },
            'colors': {
                'pdf_colors': list(self.pdf_content.get('colors', [])),
                'docx_colors': [],
                'color_accuracy': 0
            },
            'layout': {
                'pdf_blocks': len(self.pdf_content.get('text_blocks', [])),
                'docx_paragraphs': len(self.docx_content.get('paragraphs', [])),
                'structure_similarity': 0
            }
        }
        
        # Compare fonts
        pdf_fonts = set(self.pdf_content.get('fonts', []))
        docx_fonts = set(self.docx_content.get('fonts', []))
        formatting_comparison['fonts']['missing_fonts'] = list(pdf_fonts - docx_fonts)
        
        return formatting_comparison
    
    def generate_improvement_suggestions(self) -> List[Dict]:
        """Generate suggestions for improving the conversion"""
        suggestions = []
        
        text_comparison = self.compare_text_content()
        formatting_comparison = self.compare_formatting()
        
        # Text accuracy suggestions
        if text_comparison.get('similarity_ratio', 0) < 0.9:
            suggestions.append({
                'type': 'text_accuracy',
                'priority': 'high',
                'issue': f"Text similarity is only {text_comparison.get('similarity_ratio', 0):.1%}",
                'suggestion': 'Review text extraction settings and consider OCR for scanned content',
                'missing_words': text_comparison.get('missing_words', [])[:10]  # Show first 10
            })
        
        # Spacing issues
        if len(text_comparison.get('missing_words', [])) > 0:
            # Check for spacing issues
            missing_words = text_comparison.get('missing_words', [])
            potential_spacing_issues = [w for w in missing_words if len(w) > 10 and ' ' not in w]
            if potential_spacing_issues:
                suggestions.append({
                    'type': 'spacing',
                    'priority': 'medium',
                    'issue': 'Potential word spacing issues detected',
                    'suggestion': 'Add spaces between concatenated words',
                    'examples': potential_spacing_issues[:5]
                })
        
        # Font suggestions
        missing_fonts = formatting_comparison['fonts']['missing_fonts']
        if missing_fonts:
            suggestions.append({
                'type': 'fonts',
                'priority': 'medium',
                'issue': f"{len(missing_fonts)} fonts from PDF not found in DOCX",
                'suggestion': 'Install missing fonts or provide font mapping',
                'missing_fonts': missing_fonts[:5]
            })
        
        # Layout suggestions
        pdf_blocks = formatting_comparison['layout']['pdf_blocks']
        docx_paragraphs = formatting_comparison['layout']['docx_paragraphs']
        if abs(pdf_blocks - docx_paragraphs) > pdf_blocks * 0.2:  # 20% difference
            suggestions.append({
                'type': 'layout',
                'priority': 'medium',
                'issue': f"Layout structure differs significantly (PDF: {pdf_blocks} blocks, DOCX: {docx_paragraphs} paragraphs)",
                'suggestion': 'Review paragraph breaks and text block grouping'
            })
        
        return suggestions
    
    def apply_text_fixes(self, docx_path: str, output_path: str) -> Dict:
        """Apply automatic text fixes to improve accuracy"""
        try:
            doc = Document(docx_path)
            fixes_applied = []
            
            # Fix common spacing issues
            for para in doc.paragraphs:
                original_text = para.text
                
                # Fix missing spaces between words (common issue)
                # Look for patterns like "WordWord" and add spaces
                fixed_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', original_text)
                
                # Fix missing spaces after punctuation
                fixed_text = re.sub(r'([.!?])([A-Z])', r'\1 \2', fixed_text)
                
                # Fix missing spaces around common words
                fixed_text = re.sub(r'([a-z])(of|and|the|in|on|at|by)([A-Z])', r'\1 \2 \3', fixed_text)
                
                if fixed_text != original_text:
                    # Clear paragraph and rebuild with fixed text
                    para.clear()
                    para.add_run(fixed_text)
                    fixes_applied.append({
                        'type': 'spacing_fix',
                        'original': original_text[:50] + '...' if len(original_text) > 50 else original_text,
                        'fixed': fixed_text[:50] + '...' if len(fixed_text) > 50 else fixed_text
                    })
            
            # Save the fixed document
            doc.save(output_path)
            
            return {
                'status': 'success',
                'fixes_applied': len(fixes_applied),
                'details': fixes_applied,
                'output_path': output_path
            }
            
        except Exception as e:
            logger.error(f"Error applying text fixes: {e}")
            return {'status': 'error', 'error': str(e)}
    
    def full_comparison_report(self, pdf_path: str, docx_path: str) -> Dict:
        """Generate a comprehensive comparison report"""
        logger.info("Starting comprehensive PDF-DOCX comparison...")
        
        # Extract content from both files
        self.extract_pdf_content(pdf_path)
        self.extract_docx_content(docx_path)
        
        # Perform comparisons
        text_comparison = self.compare_text_content()
        formatting_comparison = self.compare_formatting()
        suggestions = self.generate_improvement_suggestions()
        
        report = {
            'pdf_file': pdf_path,
            'docx_file': docx_path,
            'text_comparison': text_comparison,
            'formatting_comparison': formatting_comparison,
            'improvement_suggestions': suggestions,
            'overall_score': self.calculate_overall_score(text_comparison, formatting_comparison),
            'timestamp': str(Path(docx_path).stat().st_mtime)
        }
        
        return report
    
    def calculate_overall_score(self, text_comp: Dict, format_comp: Dict) -> Dict:
        """Calculate an overall conversion quality score"""
        text_score = text_comp.get('similarity_ratio', 0) * 100
        
        # Font score based on how many fonts were preserved
        pdf_fonts = len(format_comp['fonts']['pdf_fonts'])
        missing_fonts = len(format_comp['fonts']['missing_fonts'])
        font_score = ((pdf_fonts - missing_fonts) / pdf_fonts * 100) if pdf_fonts > 0 else 100
        
        # Layout score based on structure similarity
        pdf_blocks = format_comp['layout']['pdf_blocks']
        docx_paras = format_comp['layout']['docx_paragraphs']
        layout_score = max(0, 100 - abs(pdf_blocks - docx_paras) / pdf_blocks * 100) if pdf_blocks > 0 else 100
        
        overall_score = (text_score * 0.6 + font_score * 0.2 + layout_score * 0.2)
        
        return {
            'overall': round(overall_score, 1),
            'text_accuracy': round(text_score, 1),
            'font_preservation': round(font_score, 1),
            'layout_similarity': round(layout_score, 1),
            'grade': self.get_grade(overall_score)
        }
    
    def get_grade(self, score: float) -> str:
        """Convert score to letter grade"""
        if score >= 90: return 'A'
        elif score >= 80: return 'B'
        elif score >= 70: return 'C'
        elif score >= 60: return 'D'
        else: return 'F'

def main():
    """Command line interface for the comparator"""
    import argparse
    
    parser = argparse.ArgumentParser(description='PDF to DOCX Comparison Tool')
    parser.add_argument('pdf_path', help='Original PDF file path')
    parser.add_argument('docx_path', help='Converted DOCX file path')
    parser.add_argument('--fix', help='Apply automatic fixes and save to this path')
    parser.add_argument('--report', help='Save detailed report to JSON file')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Create comparator
    comparator = PDFDOCXComparator()
    
    # Generate comparison report
    report = comparator.full_comparison_report(args.pdf_path, args.docx_path)
    
    # Print summary
    print(f"\n📊 PDF to DOCX Conversion Analysis")
    print(f"{'='*50}")
    print(f"PDF File: {args.pdf_path}")
    print(f"DOCX File: {args.docx_path}")
    print(f"\n🎯 Overall Score: {report['overall_score']['overall']}/100 (Grade: {report['overall_score']['grade']})")
    print(f"   Text Accuracy: {report['overall_score']['text_accuracy']}/100")
    print(f"   Font Preservation: {report['overall_score']['font_preservation']}/100")
    print(f"   Layout Similarity: {report['overall_score']['layout_similarity']}/100")
    
    print(f"\n📝 Text Comparison:")
    print(f"   Similarity: {report['text_comparison']['similarity_ratio']:.1%}")
    print(f"   PDF Words: {report['text_comparison']['word_count_pdf']}")
    print(f"   DOCX Words: {report['text_comparison']['word_count_docx']}")
    print(f"   Missing Words: {len(report['text_comparison']['missing_words'])}")
    
    print(f"\n🎨 Formatting Comparison:")
    print(f"   PDF Fonts: {len(report['formatting_comparison']['fonts']['pdf_fonts'])}")
    print(f"   DOCX Fonts: {len(report['formatting_comparison']['fonts']['docx_fonts'])}")
    print(f"   Missing Fonts: {len(report['formatting_comparison']['fonts']['missing_fonts'])}")
    
    if report['improvement_suggestions']:
        print(f"\n💡 Improvement Suggestions:")
        for i, suggestion in enumerate(report['improvement_suggestions'], 1):
            print(f"   {i}. [{suggestion['priority'].upper()}] {suggestion['issue']}")
            print(f"      → {suggestion['suggestion']}")
    
    # Apply fixes if requested
    if args.fix:
        print(f"\n🔧 Applying automatic fixes...")
        fix_result = comparator.apply_text_fixes(args.docx_path, args.fix)
        if fix_result['status'] == 'success':
            print(f"   ✅ Applied {fix_result['fixes_applied']} fixes")
            print(f"   📄 Fixed file saved to: {args.fix}")
        else:
            print(f"   ❌ Fix failed: {fix_result['error']}")
    
    # Save report if requested
    if args.report:
        with open(args.report, 'w') as f:
            json.dump(report, f, indent=2)
        print(f"\n📋 Detailed report saved to: {args.report}")

if __name__ == "__main__":
    main()
