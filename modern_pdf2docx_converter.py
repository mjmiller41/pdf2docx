#!/usr/bin/env python3
"""
Modern PDF to DOCX Converter

A completely rewritten converter using:
- pypdf for proper text and image extraction (no text-to-image conversion)
- python-docx for proper DOCX creation with template support
- Proper spacing, layout, and formatting preservation
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
import logging
import re
from io import BytesIO

# Core libraries
try:
    from pypdf import PdfReader
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.oxml.shared import OxmlElement, qn
    from PIL import Image
except ImportError as e:
    print(f"Missing required library: {e}")
    print("Install with: pip install pypdf python-docx pillow")
    sys.exit(1)

# Set up logging
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

class ModernPDF2DOCXConverter:
    """
    Modern PDF to DOCX converter that properly extracts text and applies template formatting
    """
    
    def __init__(self):
        self.pdf_reader = None
        self.docx_document = None
        self.template_document = None
        self.custom_fonts = []
        self.conversion_stats = {
            'pages_processed': 0,
            'text_blocks_extracted': 0,
            'images_extracted': 0,
            'spacing_fixes_applied': 0
        }
    
    def load_template(self, template_path: str) -> bool:
        """Load a DOCX template to apply formatting"""
        try:
            if template_path and Path(template_path).exists():
                self.template_document = Document(template_path)
                logger.info(f"Loaded template: {template_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to load template {template_path}: {e}")
            return False
    
    def register_fonts(self, font_paths: List[str]) -> None:
        """Register custom fonts for use in the document"""
        if not font_paths:
            return
        
        for font_path in font_paths:
            if Path(font_path).exists():
                font_name = Path(font_path).stem
                self.custom_fonts.append({
                    'name': font_name,
                    'path': font_path
                })
                logger.info(f"Registered font: {font_name}")
    
    def extract_pdf_content(self, pdf_path: str, password: str = None, 
                          start_page: int = 0, end_page: int = None, 
                          pages: List[int] = None) -> Dict:
        """Extract text and images from PDF using pypdf"""
        try:
            self.pdf_reader = PdfReader(pdf_path)
            
            # Handle password protection
            if self.pdf_reader.is_encrypted:
                if password:
                    self.pdf_reader.decrypt(password)
                else:
                    raise ValueError("PDF is encrypted but no password provided")
            
            total_pages = len(self.pdf_reader.pages)
            logger.info(f"PDF has {total_pages} pages")
            
            # Determine which pages to process
            if pages:
                page_indices = [p for p in pages if 0 <= p < total_pages]
            else:
                start = max(0, start_page)
                end = min(total_pages, end_page) if end_page else total_pages
                page_indices = list(range(start, end))
            
            extracted_content = {
                'pages': [],
                'total_pages': len(page_indices),
                'page_indices': page_indices
            }
            
            # Extract content from each page
            for page_idx in page_indices:
                page = self.pdf_reader.pages[page_idx]
                page_content = self._extract_page_content(page, page_idx)
                extracted_content['pages'].append(page_content)
                self.conversion_stats['pages_processed'] += 1
            
            logger.info(f"Extracted content from {len(page_indices)} pages")
            return extracted_content
            
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            return {}
    
    def _extract_page_content(self, page, page_idx: int) -> Dict:
        """Extract text and images from a single PDF page"""
        page_content = {
            'page_number': page_idx,
            'text_blocks': [],
            'images': [],
            'page_size': {
                'width': float(page.mediabox.width),
                'height': float(page.mediabox.height)
            }
        }
        
        # Extract text with positioning information
        text_blocks = self._extract_text_with_layout(page)
        page_content['text_blocks'] = text_blocks
        self.conversion_stats['text_blocks_extracted'] += len(text_blocks)
        
        # Extract images
        images = self._extract_images_from_page(page, page_idx)
        page_content['images'] = images
        self.conversion_stats['images_extracted'] += len(images)
        
        return page_content
    
    def _extract_text_with_layout(self, page) -> List[Dict]:
        """Extract text with layout information using pypdf visitor functions"""
        text_blocks = []
        current_block = None
        
        def text_visitor(text, user_matrix, tm_matrix, font_dict, font_size):
            nonlocal current_block
            
            if not text.strip():
                return
            
            # Get position information
            x, y = tm_matrix[4], tm_matrix[5]
            
            # Get font information
            font_name = font_dict.get('/BaseFont', 'Unknown') if font_dict else 'Unknown'
            if isinstance(font_name, bytes):
                font_name = font_name.decode('utf-8', errors='ignore')
            
            # Clean font name
            font_name = re.sub(r'^[A-Z]+\+', '', str(font_name))  # Remove subset prefix
            
            text_block = {
                'text': text,
                'x': float(x),
                'y': float(y),
                'font_name': font_name,
                'font_size': float(font_size) if font_size else 12.0,
                'matrix': tm_matrix
            }
            
            text_blocks.append(text_block)
        
        # Extract text with visitor
        try:
            page.extract_text(visitor_text=text_visitor)
        except Exception as e:
            logger.warning(f"Text extraction with visitor failed: {e}, falling back to simple extraction")
            # Fallback to simple text extraction
            simple_text = page.extract_text()
            if simple_text.strip():
                text_blocks.append({
                    'text': simple_text,
                    'x': 0,
                    'y': 0,
                    'font_name': 'Unknown',
                    'font_size': 12.0,
                    'matrix': [1, 0, 0, 1, 0, 0]
                })
        
        # Sort text blocks by position (top to bottom, left to right)
        text_blocks.sort(key=lambda b: (-b['y'], b['x']))
        
        return text_blocks
    
    def _extract_images_from_page(self, page, page_idx: int) -> List[Dict]:
        """Extract images from a PDF page"""
        images = []
        
        try:
            for img_idx, image_file_object in enumerate(page.images):
                try:
                    # Get image data
                    image_data = image_file_object.data
                    image_name = image_file_object.name or f"image_{page_idx}_{img_idx}"
                    
                    # Try to get image dimensions
                    try:
                        pil_image = Image.open(BytesIO(image_data))
                        width, height = pil_image.size
                        format_type = pil_image.format
                    except Exception:
                        width, height = 100, 100  # Default size
                        format_type = 'Unknown'
                    
                    image_info = {
                        'name': image_name,
                        'data': image_data,
                        'width': width,
                        'height': height,
                        'format': format_type,
                        'page_number': page_idx,
                        'index': img_idx
                    }
                    
                    images.append(image_info)
                    
                except Exception as e:
                    logger.warning(f"Failed to extract image {img_idx} from page {page_idx}: {e}")
                    
        except Exception as e:
            logger.warning(f"Failed to extract images from page {page_idx}: {e}")
        
        return images
    
    def create_docx_document(self, extracted_content: Dict, template_path: str = None) -> Document:
        """Create a DOCX document from extracted PDF content"""
        
        # Load template or create new document
        if template_path and self.load_template(template_path):
            # Create new document based on template
            self.docx_document = Document(template_path)
            # Clear existing content but keep styles and formatting
            for paragraph in self.docx_document.paragraphs[:]:
                p = paragraph._element
                p.getparent().remove(p)
        else:
            # Create new document
            self.docx_document = Document()
        
        # Apply template formatting if available
        self._apply_template_formatting(extracted_content)
        
        # Process each page
        for page_content in extracted_content['pages']:
            self._process_page_content(page_content)
            
            # Add page break between pages (except for the last page)
            if page_content != extracted_content['pages'][-1]:
                self.docx_document.add_page_break()
        
        return self.docx_document
    
    def _apply_template_formatting(self, extracted_content: Dict) -> None:
        """Apply template formatting to the document"""
        if not self.template_document:
            # Set default formatting for new documents
            self._set_default_formatting(extracted_content)
            return
        
        # Copy template formatting
        try:
            # Get template section
            template_section = self.template_document.sections[0]
            doc_section = self.docx_document.sections[0]
            
            # Copy page setup
            doc_section.page_width = template_section.page_width
            doc_section.page_height = template_section.page_height
            doc_section.orientation = template_section.orientation
            
            # Copy margins
            doc_section.left_margin = template_section.left_margin
            doc_section.right_margin = template_section.right_margin
            doc_section.top_margin = template_section.top_margin
            doc_section.bottom_margin = template_section.bottom_margin
            
            logger.info("Applied template formatting to document")
            
        except Exception as e:
            logger.warning(f"Failed to apply template formatting: {e}")
            self._set_default_formatting(extracted_content)
    
    def _set_default_formatting(self, extracted_content: Dict) -> None:
        """Set default formatting based on PDF page size"""
        try:
            section = self.docx_document.sections[0]
            
            # Get PDF page dimensions (use first page)
            if extracted_content['pages']:
                pdf_width = extracted_content['pages'][0]['page_size']['width']
                pdf_height = extracted_content['pages'][0]['page_size']['height']
                
                # Convert PDF points to inches (72 points = 1 inch)
                page_width_inches = pdf_width / 72
                page_height_inches = pdf_height / 72
                
                # Set page size
                section.page_width = Inches(page_width_inches)
                section.page_height = Inches(page_height_inches)
                
                # Determine orientation
                if pdf_width > pdf_height:
                    section.orientation = WD_ORIENT.LANDSCAPE
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                
                # Set reasonable margins (1 inch default)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                
                logger.info(f"Set page size: {page_width_inches:.1f}\" x {page_height_inches:.1f}\"")
            
        except Exception as e:
            logger.warning(f"Failed to set default formatting: {e}")
    
    def _process_page_content(self, page_content: Dict) -> None:
        """Process content from a single PDF page"""
        
        # Group text blocks into paragraphs based on position
        paragraphs = self._group_text_into_paragraphs(page_content['text_blocks'])
        
        # Add text paragraphs to document
        for paragraph_data in paragraphs:
            self._add_paragraph_to_document(paragraph_data)
        
        # Add images to document
        for image_data in page_content['images']:
            self._add_image_to_document(image_data)
    
    def _group_text_into_paragraphs(self, text_blocks: List[Dict]) -> List[Dict]:
        """Group text blocks into logical paragraphs"""
        if not text_blocks:
            return []
        
        paragraphs = []
        current_paragraph = {
            'text_parts': [],
            'y_position': None,
            'font_info': {}
        }
        
        # Group text blocks by approximate Y position
        y_threshold = 5  # Points threshold for same line
        
        for block in text_blocks:
            if current_paragraph['y_position'] is None:
                # First block
                current_paragraph['y_position'] = block['y']
                current_paragraph['text_parts'].append(block)
            elif abs(block['y'] - current_paragraph['y_position']) <= y_threshold:
                # Same line
                current_paragraph['text_parts'].append(block)
            else:
                # New line/paragraph
                if current_paragraph['text_parts']:
                    paragraphs.append(current_paragraph)
                
                current_paragraph = {
                    'text_parts': [block],
                    'y_position': block['y'],
                    'font_info': {}
                }
        
        # Add the last paragraph
        if current_paragraph['text_parts']:
            paragraphs.append(current_paragraph)
        
        return paragraphs
    
    def _add_paragraph_to_document(self, paragraph_data: Dict) -> None:
        """Add a paragraph to the DOCX document with proper formatting"""
        
        # Combine text parts with proper spacing
        combined_text = self._combine_text_with_spacing(paragraph_data['text_parts'])
        
        if not combined_text.strip():
            return
        
        # Create paragraph
        paragraph = self.docx_document.add_paragraph()
        
        # Add text runs with formatting
        for text_part in paragraph_data['text_parts']:
            run = paragraph.add_run(text_part['text'])
            
            # Apply font formatting
            self._apply_font_formatting(run, text_part)
        
        # Apply paragraph formatting
        self._apply_paragraph_formatting(paragraph, paragraph_data)
    
    def _combine_text_with_spacing(self, text_parts: List[Dict]) -> str:
        """Combine text parts with proper spacing"""
        if not text_parts:
            return ""
        
        combined_parts = []
        
        for i, part in enumerate(text_parts):
            text = part['text']
            
            # Fix common spacing issues
            text = self._fix_text_spacing(text)
            
            # Add space between parts if needed
            if i > 0:
                prev_part = text_parts[i-1]
                # Calculate distance between text parts
                x_distance = abs(part['x'] - (prev_part['x'] + len(prev_part['text']) * 5))  # Rough estimate
                
                # Add space if parts are separated
                if x_distance > 10 and not text.startswith(' ') and not combined_parts[-1].endswith(' '):
                    combined_parts.append(' ')
            
            combined_parts.append(text)
        
        combined_text = ''.join(combined_parts)
        
        # Apply additional spacing fixes
        combined_text = self._fix_text_spacing(combined_text)
        
        return combined_text
    
    def _fix_text_spacing(self, text: str) -> str:
        """Fix common text spacing issues"""
        original_text = text
        
        # Fix missing spaces between words (common in PDF extraction)
        # Pattern: lowercase letter followed by uppercase letter
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Fix missing spaces after punctuation
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)
        
        # Fix missing spaces around common words
        text = re.sub(r'([a-z])(of|and|the|in|on|at|by|for|with|to)([A-Z])', r'\1 \2 \3', text)
        
        # Fix missing spaces before numbers
        text = re.sub(r'([a-z])(\d)', r'\1 \2', text)
        
        # Fix multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        if text != original_text:
            self.conversion_stats['spacing_fixes_applied'] += 1
        
        return text
    
    def _apply_font_formatting(self, run, text_part: Dict) -> None:
        """Apply font formatting to a text run"""
        try:
            font = run.font
            
            # Set font name
            font_name = text_part.get('font_name', 'Calibri')
            
            # Check if we have a custom font mapping
            custom_font = self._get_custom_font(font_name)
            if custom_font:
                font.name = custom_font['name']
            else:
                # Clean and set font name
                clean_font_name = self._clean_font_name(font_name)
                font.name = clean_font_name
            
            # Set font size
            font_size = text_part.get('font_size', 12.0)
            font.size = Pt(max(8, min(72, font_size)))  # Clamp between 8 and 72 pt
            
        except Exception as e:
            logger.warning(f"Failed to apply font formatting: {e}")
    
    def _get_custom_font(self, font_name: str) -> Optional[Dict]:
        """Get custom font mapping if available"""
        for custom_font in self.custom_fonts:
            if font_name.lower() in custom_font['name'].lower():
                return custom_font
        return None
    
    def _clean_font_name(self, font_name: str) -> str:
        """Clean and normalize font name"""
        if not font_name or font_name == 'Unknown':
            return 'Calibri'
        
        # Remove common prefixes and suffixes
        font_name = re.sub(r'^[A-Z]+\+', '', font_name)  # Remove subset prefix
        font_name = re.sub(r'[,\-].*$', '', font_name)   # Remove variants
        
        # Map common PDF fonts to Word fonts
        font_mapping = {
            'Times': 'Times New Roman',
            'Helvetica': 'Arial',
            'Courier': 'Courier New',
            'Symbol': 'Symbol',
            'ZapfDingbats': 'Wingdings'
        }
        
        for pdf_font, word_font in font_mapping.items():
            if pdf_font.lower() in font_name.lower():
                return word_font
        
        return font_name if font_name else 'Calibri'
    
    def _apply_paragraph_formatting(self, paragraph, paragraph_data: Dict) -> None:
        """Apply paragraph-level formatting"""
        try:
            paragraph_format = paragraph.paragraph_format
            
            # Set line spacing
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # Set spacing before/after
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(6)
            
        except Exception as e:
            logger.warning(f"Failed to apply paragraph formatting: {e}")
    
    def _add_image_to_document(self, image_data: Dict) -> None:
        """Add an image to the DOCX document"""
        try:
            # Create image from data
            image_stream = BytesIO(image_data['data'])
            
            # Calculate appropriate size (fit within page margins)
            section = self.docx_document.sections[0]
            max_width = section.page_width - section.left_margin - section.right_margin
            max_height = section.page_height - section.top_margin - section.bottom_margin
            
            # Get original image dimensions
            original_width = image_data['width']
            original_height = image_data['height']
            
            # Calculate scaling to fit within margins
            max_width_inches = max_width.inches
            max_height_inches = max_height.inches
            
            # Convert pixels to inches (assuming 72 DPI)
            original_width_inches = original_width / 72
            original_height_inches = original_height / 72
            
            width_scale = max_width_inches / original_width_inches
            height_scale = max_height_inches / original_height_inches
            scale = min(width_scale, height_scale, 1.0)  # Don't upscale
            
            # Calculate final dimensions
            final_width = Inches(original_width_inches * scale)
            final_height = Inches(original_height_inches * scale)
            
            # Add image to document
            paragraph = self.docx_document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_stream, width=final_width, height=final_height)
            
            # Center the image
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            logger.info(f"Added image: {image_data['name']} ({final_width.inches:.1f}\" x {final_height.inches:.1f}\")")
            
        except Exception as e:
            logger.warning(f"Failed to add image {image_data['name']}: {e}")
    
    def convert_pdf_to_docx(self, pdf_path: str, output_path: str, 
                          template_path: str = None, font_paths: List[str] = None,
                          password: str = None, start_page: int = 0, 
                          end_page: int = None, pages: List[int] = None) -> Dict:
        """
        Convert PDF to DOCX with proper text extraction and template formatting
        
        Args:
            pdf_path: Path to input PDF file
            output_path: Path for output DOCX file
            template_path: Optional path to DOCX template
            font_paths: Optional list of custom font files
            password: Optional PDF password
            start_page: Starting page (0-indexed)
            end_page: Ending page (exclusive)
            pages: Specific pages to convert
            
        Returns:
            Dict with conversion results
        """
        
        try:
            logger.info(f"Starting conversion: {pdf_path} -> {output_path}")
            
            # Reset stats
            self.conversion_stats = {
                'pages_processed': 0,
                'text_blocks_extracted': 0,
                'images_extracted': 0,
                'spacing_fixes_applied': 0
            }
            
            # Register custom fonts
            if font_paths:
                self.register_fonts(font_paths)
            
            # Extract content from PDF
            extracted_content = self.extract_pdf_content(
                pdf_path, password, start_page, end_page, pages
            )
            
            if not extracted_content:
                return {'status': 'error', 'error': 'Failed to extract PDF content'}
            
            # Create DOCX document
            document = self.create_docx_document(extracted_content, template_path)
            
            # Save document
            document.save(output_path)
            
            # Prepare result
            result = {
                'status': 'success',
                'output_path': output_path,
                'stats': self.conversion_stats,
                'pages_converted': extracted_content['total_pages']
            }
            
            logger.info(f"Conversion completed successfully!")
            logger.info(f"Stats: {self.conversion_stats}")
            
            return result
            
        except Exception as e:
            error_msg = f"Conversion failed: {str(e)}"
            logger.error(error_msg)
            return {'status': 'error', 'error': error_msg}

def main():
    """Command line interface"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Modern PDF to DOCX Converter')
    parser.add_argument('pdf_path', help='Input PDF file path')
    parser.add_argument('output_path', help='Output DOCX file path')
    parser.add_argument('--template', help='DOCX template file path')
    parser.add_argument('--fonts', nargs='+', help='Custom font file paths')
    parser.add_argument('--password', help='PDF password')
    parser.add_argument('--start-page', type=int, default=0, help='Start page (0-indexed)')
    parser.add_argument('--end-page', type=int, help='End page (exclusive)')
    parser.add_argument('--pages', nargs='+', type=int, help='Specific pages to convert')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Create converter
    converter = ModernPDF2DOCXConverter()
    
    # Perform conversion
    result = converter.convert_pdf_to_docx(
        pdf_path=args.pdf_path,
        output_path=args.output_path,
        template_path=args.template,
        font_paths=args.fonts,
        password=args.password,
        start_page=args.start_page,
        end_page=args.end_page,
        pages=args.pages
    )
    
    # Print results
    if result['status'] == 'success':
        print(f"✅ Conversion successful!")
        print(f"📄 Output: {result['output_path']}")
        print(f"📊 Pages converted: {result['pages_converted']}")
        print(f"📈 Stats: {result['stats']}")
    else:
        print(f"❌ Conversion failed: {result['error']}")
        sys.exit(1)

if __name__ == "__main__":
    main()
