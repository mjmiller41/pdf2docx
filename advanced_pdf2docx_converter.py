#!/usr/bin/env python3
"""
Advanced PDF to DOCX Converter

This converter properly extracts text and images from PDFs and formats them
according to template margins and styles, avoiding full-page image conversion.
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
from PIL import Image
import os
import sys
from pathlib import Path
import tempfile
import logging
from typing import List, Dict, Tuple, Optional
import re

# Set up logging
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

class AdvancedPDF2DOCXConverter:
    """Advanced PDF to DOCX converter with proper text extraction and template support"""
    
    def __init__(self):
        self.font_registry = {}
        self.template_doc = None
        self.default_margins = {
            'top': Inches(1),
            'bottom': Inches(1),
            'left': Inches(1),
            'right': Inches(1)
        }
    
    def register_fonts(self, font_paths: List[str]) -> int:
        """Register custom fonts for use in the document"""
        font_count = 0
        for font_path in font_paths:
            if os.path.isdir(font_path):
                # Register all fonts in directory
                for font_file in Path(font_path).glob('*.ttf'):
                    self.font_registry[font_file.stem] = str(font_file)
                    font_count += 1
                for font_file in Path(font_path).glob('*.otf'):
                    self.font_registry[font_file.stem] = str(font_file)
                    font_count += 1
            elif os.path.isfile(font_path) and font_path.lower().endswith(('.ttf', '.otf')):
                font_name = Path(font_path).stem
                self.font_registry[font_name] = font_path
                font_count += 1
        
        logger.info(f"Registered {font_count} custom fonts")
        return font_count
    
    def load_template(self, template_path: str) -> bool:
        """Load a Word template (.dotx or .docx) for styling"""
        try:
            self.template_doc = Document(template_path)
            logger.info(f"Loaded template: {template_path}")
            
            # Extract margins from template
            section = self.template_doc.sections[0]
            self.default_margins = {
                'top': section.top_margin,
                'bottom': section.bottom_margin,
                'left': section.left_margin,
                'right': section.right_margin
            }
            logger.info(f"Template margins: {self.default_margins}")
            return True
        except Exception as e:
            logger.warning(f"Could not load template: {e}")
            return False
    
    def extract_text_blocks(self, page) -> List[Dict]:
        """Extract text blocks with formatting information"""
        blocks = []
        text_dict = page.get_text("dict")
        
        for block in text_dict["blocks"]:
            if "lines" in block:  # Text block
                for line in block["lines"]:
                    for span in line["spans"]:
                        if span["text"].strip():
                            blocks.append({
                                'type': 'text',
                                'text': span["text"],
                                'bbox': span["bbox"],
                                'font': span["font"],
                                'size': span["size"],
                                'flags': span["flags"],
                                'color': span["color"]
                            })
        
        return blocks
    
    def extract_images(self, page, page_num: int) -> List[Dict]:
        """Extract images from page and prepare them for insertion"""
        images = []
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            try:
                # Get image data
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    
                    # Get image position on page
                    img_rect = page.get_image_rects(xref)
                    if img_rect:
                        bbox = img_rect[0]
                        
                        # Process image to fit within margins
                        processed_img = self.process_image(img_data, bbox, page.rect)
                        
                        if processed_img:
                            images.append({
                                'type': 'image',
                                'data': processed_img,
                                'bbox': bbox,
                                'original_size': (pix.width, pix.height)
                            })
                
                pix = None  # Free memory
                
            except Exception as e:
                logger.warning(f"Could not extract image {img_index} from page {page_num}: {e}")
        
        return images
    
    def process_image(self, img_data: bytes, bbox: Tuple, page_rect: fitz.Rect) -> Optional[bytes]:
        """Process image to fit within document margins"""
        try:
            # Open image with PIL
            img = Image.open(io.BytesIO(img_data))
            
            # Calculate maximum size based on page margins
            page_width = page_rect.width
            page_height = page_rect.height
            
            # Convert margins to points (assuming 72 DPI)
            margin_left = self.default_margins['left'].pt
            margin_right = self.default_margins['right'].pt
            margin_top = self.default_margins['top'].pt
            margin_bottom = self.default_margins['bottom'].pt
            
            max_width = page_width - margin_left - margin_right
            max_height = page_height - margin_top - margin_bottom
            
            # Resize image to fit within margins while maintaining aspect ratio
            img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            # Convert back to bytes
            output = io.BytesIO()
            img.save(output, format='PNG')
            return output.getvalue()
            
        except Exception as e:
            logger.warning(f"Could not process image: {e}")
            return None
    
    def apply_template_style(self, doc: Document) -> None:
        """Apply template styles to the document"""
        if not self.template_doc:
            return
        
        try:
            # Copy styles from template
            for style in self.template_doc.styles:
                if style.name not in [s.name for s in doc.styles]:
                    # This is a simplified approach - in practice, style copying is complex
                    logger.debug(f"Would copy style: {style.name}")
            
            # Apply template margins to all sections
            for section in doc.sections:
                section.top_margin = self.default_margins['top']
                section.bottom_margin = self.default_margins['bottom']
                section.left_margin = self.default_margins['left']
                section.right_margin = self.default_margins['right']
                
        except Exception as e:
            logger.warning(f"Could not fully apply template styles: {e}")
    
    def add_text_to_document(self, doc: Document, text_blocks: List[Dict]) -> None:
        """Add extracted text to document with proper formatting"""
        if not text_blocks:
            return
        
        # Group text blocks by approximate line (similar Y coordinates)
        lines = {}
        for block in text_blocks:
            y_coord = int(block['bbox'][1])  # Top Y coordinate
            if y_coord not in lines:
                lines[y_coord] = []
            lines[y_coord].append(block)
        
        # Sort lines by Y coordinate (top to bottom)
        sorted_lines = sorted(lines.items())
        
        for y_coord, line_blocks in sorted_lines:
            # Sort blocks in line by X coordinate (left to right)
            line_blocks.sort(key=lambda b: b['bbox'][0])
            
            # Create paragraph for this line
            paragraph = doc.add_paragraph()
            
            for block in line_blocks:
                run = paragraph.add_run(block['text'])
                
                # Apply font formatting
                font_name = block['font']
                font_size = block['size']
                
                # Use custom font if available
                if self.font_registry:
                    # Try to match font name with registered fonts
                    for reg_font_name, reg_font_path in self.font_registry.items():
                        if reg_font_name.lower() in font_name.lower():
                            run.font.name = reg_font_name
                            break
                    else:
                        run.font.name = font_name
                else:
                    run.font.name = font_name
                
                run.font.size = Pt(font_size)
                
                # Apply text formatting based on flags
                if block['flags'] & 2**4:  # Bold
                    run.font.bold = True
                if block['flags'] & 2**1:  # Italic
                    run.font.italic = True
                
                # Apply color if not black
                if block['color'] != 0:
                    # Convert color from integer to RGB
                    r = (block['color'] >> 16) & 255
                    g = (block['color'] >> 8) & 255
                    b = block['color'] & 255
                    run.font.color.rgb = RGBColor(r, g, b)
    
    def add_images_to_document(self, doc: Document, images: List[Dict]) -> None:
        """Add processed images to document"""
        for img_data in images:
            try:
                # Add image to document
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                
                # Add image from bytes
                img_stream = io.BytesIO(img_data['data'])
                run.add_picture(img_stream)
                
                # Center the image
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            except Exception as e:
                logger.warning(f"Could not add image to document: {e}")
    
    def convert_pdf_to_docx(self, pdf_path: str, output_path: str, 
                           template_path: Optional[str] = None,
                           font_paths: Optional[List[str]] = None,
                           start_page: int = 0, end_page: Optional[int] = None,
                           pages: Optional[List[int]] = None,
                           password: Optional[str] = None) -> Dict:
        """
        Convert PDF to DOCX with proper text extraction and template formatting
        """
        try:
            logger.info(f"Starting advanced conversion: {pdf_path} -> {output_path}")
            
            # Register fonts if provided
            if font_paths:
                self.register_fonts(font_paths)
            
            # Load template if provided
            if template_path:
                self.load_template(template_path)
            
            # Open PDF
            pdf_doc = fitz.open(pdf_path)
            
            # Handle password
            if password:
                if not pdf_doc.authenticate(password):
                    return {'status': 'error', 'error': 'Invalid password'}
            
            # Create new Word document
            if self.template_doc:
                # Start with template
                doc = Document(template_path)
                # Clear existing content
                for element in doc.element.body:
                    doc.element.body.remove(element)
            else:
                doc = Document()
            
            # Apply template styles and margins
            self.apply_template_style(doc)
            
            # Determine pages to process
            total_pages = len(pdf_doc)
            if pages:
                page_list = [p for p in pages if 0 <= p < total_pages]
            else:
                start = max(0, start_page)
                end = min(total_pages, end_page) if end_page else total_pages
                page_list = list(range(start, end))
            
            logger.info(f"Processing {len(page_list)} pages")
            
            # Process each page
            for page_num in page_list:
                logger.info(f"Processing page {page_num + 1}/{total_pages}")
                
                page = pdf_doc[page_num]
                
                # Extract text blocks
                text_blocks = self.extract_text_blocks(page)
                
                # Extract and process images
                images = self.extract_images(page, page_num)
                
                # Add page break if not first page
                if page_num > page_list[0]:
                    doc.add_page_break()
                
                # Add text to document
                if text_blocks:
                    self.add_text_to_document(doc, text_blocks)
                    logger.info(f"Added {len(text_blocks)} text blocks from page {page_num + 1}")
                else:
                    logger.warning(f"No text found on page {page_num + 1}")
                
                # Add images to document
                if images:
                    self.add_images_to_document(doc, images)
                    logger.info(f"Added {len(images)} images from page {page_num + 1}")
            
            # Save document
            doc.save(output_path)
            pdf_doc.close()
            
            # Verify output
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                file_size = os.path.getsize(output_path)
                logger.info(f"Conversion successful: {output_path} ({file_size} bytes)")
                
                return {
                    'status': 'success',
                    'output_path': output_path,
                    'file_size': file_size,
                    'pages_processed': len(page_list),
                    'fonts_used': len(self.font_registry),
                    'template_applied': template_path is not None
                }
            else:
                return {'status': 'error', 'error': 'Output file was not created or is empty'}
                
        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            return {'status': 'error', 'error': str(e)}

def main():
    """Command line interface for the advanced converter"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Advanced PDF to DOCX Converter')
    parser.add_argument('pdf_path', help='Input PDF file path')
    parser.add_argument('output_path', help='Output DOCX file path')
    parser.add_argument('--template', help='Word template file (.dotx or .docx)')
    parser.add_argument('--fonts', action='append', help='Font files or directories (can be used multiple times)')
    parser.add_argument('--start-page', type=int, default=0, help='Start page (0-indexed)')
    parser.add_argument('--end-page', type=int, help='End page (exclusive)')
    parser.add_argument('--pages', help='Specific pages (comma-separated, 0-indexed)')
    parser.add_argument('--password', help='PDF password')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Parse pages if provided
    pages_list = None
    if args.pages:
        try:
            pages_list = [int(p.strip()) for p in args.pages.split(',')]
        except ValueError:
            print("Error: Invalid pages format. Use comma-separated integers.")
            sys.exit(1)
    
    # Create converter
    converter = AdvancedPDF2DOCXConverter()
    
    # Perform conversion
    result = converter.convert_pdf_to_docx(
        pdf_path=args.pdf_path,
        output_path=args.output_path,
        template_path=args.template,
        font_paths=args.fonts,
        start_page=args.start_page,
        end_page=args.end_page,
        pages=pages_list,
        password=args.password
    )
    
    if result['status'] == 'success':
        print(f"✅ Conversion successful!")
        print(f"   Output: {result['output_path']}")
        print(f"   Size: {result['file_size']} bytes")
        print(f"   Pages: {result['pages_processed']}")
        if result['fonts_used']:
            print(f"   Custom fonts: {result['fonts_used']}")
        if result['template_applied']:
            print(f"   Template applied: ✅")
    else:
        print(f"❌ Conversion failed: {result['error']}")
        sys.exit(1)

if __name__ == "__main__":
    main()
